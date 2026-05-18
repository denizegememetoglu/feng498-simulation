"""Generate the policy-comparison Word doc from output/policy_stats.json.

Run from repo root:
    python tools/build_analysis_doc.py

Output: /home/dege/Downloads/feng498-policy-analysis.docx

Interpretations are derived from the data (ANOVA p, winner, Tukey counts),
not hard-coded — re-running after a sim refresh always produces a doc
consistent with the latest numbers.
"""

import json
import math
import os
import sys
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm

IN = os.path.join(os.path.dirname(__file__), "..", "output", "policy_stats.json")
OUT = "/home/dege/Downloads/feng498-policy-analysis.docx"


def h1(d, t): return d.add_heading(t, level=1)
def h2(d, t): return d.add_heading(t, level=2)
def h3(d, t): return d.add_heading(t, level=3)


def para(d, text, bold=False, italic=False):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    return p


def bullet(d, t): return d.add_paragraph(t, style="List Bullet")


def code(d, t):
    p = d.add_paragraph()
    r = p.add_run(t)
    r.font.name, r.font.size = "Consolas", Pt(9)
    return p


def table(d, rows, headers):
    t = d.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.rows[i].cells[j].text = str(v)
    return t


def fmt(x, p=3):
    if x is None: return "n/a"
    if isinstance(x, float):
        if x != x: return "n/a"
        return f"{x:.{p}f}"
    return str(x)


def fmt_p(p):
    if p is None: return "n/a"
    return "<0.0001" if p < 0.0001 else f"{p:.4f}"


def effect_mag(d):
    if d is None or d != d: return "n/a"
    a = abs(d)
    return ("trivial" if a < 0.2 else "small" if a < 0.5 else
            "medium" if a < 0.8 else "large" if a < 1.5 else "very large")


def derived_interpretation(block):
    """Build a 1-paragraph interpretation directly from the data."""
    a = block["anova"]
    pol = block["per_policy"]
    sig = a and a["p_value"] < 0.05
    winner = block["winner"]
    win_mean = pol[winner]["mean"]
    direction = block["direction"]

    tukey = block.get("tukey_hsd") or []
    n_sig = sum(1 for t in tukey if t["significant_at_0.05"])

    welch = block.get("welch_vs_baseline") or []
    largest_d = max(welch, key=lambda w: abs(w["cohens_d"] or 0),
                    default=None) if welch else None

    cvs = [s["cv"] for s in pol.values() if s["cv"] == s["cv"]]
    max_cv = max(cvs) if cvs else 0

    parts = []
    if a:
        parts.append(
            f"ANOVA F({a['df_between']},{a['df_within']}) = {a['F']:.2f}, "
            f"p = {fmt_p(a['p_value'])} — "
            f"{'reject H0, at least one policy mean differs' if sig else 'fail to reject H0, no detectable difference between policies'}."
        )
    parts.append(
        f"Best mean: {winner} at {fmt(win_mean)} "
        f"({direction} is better)."
    )
    if tukey:
        parts.append(
            f"Tukey HSD: {n_sig} of {len(tukey)} pairs significant at α=0.05."
        )
    if largest_d and largest_d["cohens_d"] is not None:
        parts.append(
            f"Largest Welch effect vs Heuristic: {largest_d['policy']} "
            f"(d = {fmt(largest_d['cohens_d'])}, "
            f"{effect_mag(largest_d['cohens_d'])})."
        )
    if max_cv > 0.15:
        parts.append(
            f"Note: max CV across policies = {fmt(max_cv)} — high variance, "
            f"detection power is limited; bump N_REPLICATIONS to tighten CIs."
        )
    return " ".join(parts)


def per_policy_table(d, block):
    rows = []
    for pol, st in block["per_policy"].items():
        rows.append([pol, fmt(st["mean"]),
                     f"± {fmt(st['ci95_half_width'])}",
                     fmt(st["stdev"]), fmt(st["cv"]),
                     fmt(st["min"]), fmt(st["max"])])
    table(d, rows, ["Policy", "Mean", "95% CI half-width",
                    "Stdev", "CV", "Min", "Max"])


def tukey_table(d, block):
    pairs = block.get("tukey_hsd") or []
    if not pairs: return
    rows = []
    for t in pairs:
        rows.append([
            "*" if t["significant_at_0.05"] else "",
            t["policy_a"], t["policy_b"], fmt(t["mean_diff"]),
            f"[{fmt(t['ci_low'])}, {fmt(t['ci_high'])}]",
            fmt_p(t["p_value"]),
        ])
    table(d, rows, ["sig", "Policy A", "Policy B",
                    "Mean diff (A−B)", "95% CI", "p"])


def welch_table(d, block):
    welch = block.get("welch_vs_baseline") or []
    if not welch: return
    rows = []
    for w in welch:
        rows.append([
            "*" if w["significant_at_0.05"] else "",
            w["policy"], fmt(w["t"]), fmt_p(w["p_value"]),
            fmt(w["cohens_d"]), effect_mag(w["cohens_d"]),
        ])
    table(d, rows, ["sig", "Policy (vs Heuristic)", "t",
                    "p", "Cohen's d", "Effect size"])


def main():
    with open(IN) as f:
        report = json.load(f)

    n_reps = next(iter(report["n_replications"].values()), 0)
    n_pol = len(report["n_replications"])
    metrics = report["metrics"]

    sig_count = sum(1 for b in metrics.values()
                    if b["anova"] and b["anova"]["p_value"] < 0.05)

    d = Document()
    for s in d.sections:
        s.left_margin = s.right_margin = Cm(2.0)
        s.top_margin = s.bottom_margin = Cm(2.0)

    # Title
    t = d.add_paragraph()
    r = t.add_run("FENG 498 — Policy Comparison Statistical Analysis")
    r.bold = True
    r.font.size = Pt(20)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = d.add_paragraph()
    r = sub.add_run(
        f"Schneider Manisa warehouse simulation · Deniz Ege Memetoğlu\n"
        f"{n_pol} policies × {n_reps} replications · auto-generated from "
        f"output/policy_stats.json"
    )
    r.italic = True
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    d.add_paragraph()

    # 1. Summary
    h1(d, "1. Executive summary")
    para(d,
         f"Compared {n_pol} slotting policies across {n_reps} independent "
         f"replications each (different random seeds, same SimPy model). "
         f"For every KPI: one-way ANOVA across policies, Tukey HSD pairwise "
         f"post-hoc, pairwise Welch's t-tests against the Heuristic baseline "
         f"with Cohen's d effect sizes.")
    bullet(d,
           f"{sig_count} of {len(metrics)} KPIs show a statistically "
           f"significant policy effect at α = 0.05 (ANOVA p < 0.05).")
    for key, block in metrics.items():
        a = block["anova"]
        if a and a["p_value"] < 0.05:
            bullet(d,
                   f"{block['label']}: winner {block['winner']} "
                   f"(ANOVA p = {fmt_p(a['p_value'])}, "
                   f"{block['direction']} is better).")

    # 2. Methodology
    h1(d, "2. Methodology")
    bullet(d, "One-way ANOVA across policies (scipy.stats.f_oneway).")
    bullet(d, "Tukey HSD pairwise post-hoc (scipy.stats.tukey_hsd) — controls family-wise error.")
    bullet(d, "Welch's t-test of each non-baseline policy vs Heuristic (scipy.stats.ttest_ind, equal_var=False).")
    bullet(d, "Cohen's d (pooled stdev) — 0.2 small, 0.5 medium, 0.8 large.")
    bullet(d, f"95% Student-t CI per policy (df = N − 1 = {n_reps - 1}).")
    bullet(d, "Coefficient of variation (CV) per policy as a replication-adequacy check.")

    para(d,
         "Why ANOVA (not chi-square) for policy comparison: chi-square in "
         "src/validate.py answers whether simulation marginal distribution "
         "matches SAP. Here the question is whether continuous-KPI means "
         "differ between k > 2 experimental groups (policies) given "
         "replication-level noise — ANOVA is the standard test.",
         italic=True)

    # 3. Per-KPI results
    h1(d, "3. Per-KPI results")
    order_priority = ["avg_walk_distance", "total_walk_distance",
                      "reach_truck_utilization", "avg_prep_time",
                      "orders_completed", "avg_lead_time",
                      "avg_op_queue_wait", "operator_utilization"]
    seen = set()
    for k in order_priority + list(metrics.keys()):
        if k in seen or k not in metrics: continue
        seen.add(k)
        block = metrics[k]
        a = block["anova"]
        h2(d, block["label"])
        table(d, [
            ("Direction", f"{block['direction']} is better"),
            ("Winner", block["winner"]),
            ("ANOVA F", fmt(a["F"]) if a else "n/a"),
            ("ANOVA p", fmt_p(a["p_value"]) if a else "n/a"),
            ("Verdict (α=0.05)",
             "REJECT H0 — means differ" if a and a["p_value"] < 0.05
             else "Fail to reject H0"),
        ], ["Item", "Value"])
        h3(d, "Per-policy descriptives")
        per_policy_table(d, block)
        h3(d, "Tukey HSD pairwise")
        para(d, "* in 'sig' column = significant at α=0.05.", italic=True)
        tukey_table(d, block)
        h3(d, "Welch's t-test vs Heuristic")
        welch_table(d, block)
        h3(d, "Interpretation")
        para(d, derived_interpretation(block))
        d.add_paragraph()

    # 4. Replication adequacy
    h1(d, "4. Replication adequacy")
    para(d, "Maximum CV per KPI across policies (rule of thumb: < 0.10 means N is enough):")
    rows = []
    for key, block in metrics.items():
        cvs = [(p, st["cv"]) for p, st in block["per_policy"].items()]
        worst = max(cvs, key=lambda x: x[1] if x[1] == x[1] else 0)
        rows.append([block["label"], worst[0], fmt(worst[1])])
    table(d, rows, ["KPI", "Worst-CV policy", "CV"])

    # 5. Reproducibility
    h1(d, "5. How to reproduce")
    code(d,
         "python -m src.main               # writes output/replications.json\n"
         "python -m src.analyze            # writes output/policy_stats.json\n"
         "python tools/build_analysis_doc.py  # writes this Word doc\n"
         "cp output/policy_stats.json docs/  # refreshes GH Pages dashboard\n")

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    d.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
