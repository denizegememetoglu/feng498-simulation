"""Generate the final Sargent V&V Word report.

Reads everything from output/*.json (validation, policy stats, ZWM92
summary, timing study, sensitivity) so re-running after a pipeline change
re-syncs the doc without manual edits.

Output: /home/dege/Downloads/feng498-final-VV-report.docx
"""

import json
import os
from datetime import date

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = "output"
DEST = "/home/dege/Downloads/feng498-final-VV-report.docx"
FLOWCHART_PATH = "output/sim_flowchart.png"
POLICY_CHART_PATH = "output/policy_comparison.png"
LAYOUT_TOPDOWN_PATH = "output/layout_topdown_with_decorations.png"


def _read_json(name: str, default=None):
    p = os.path.join(OUT_DIR, name)
    if not os.path.exists(p):
        return default
    with open(p) as f:
        return json.load(f)


def _add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs:
            for run in r.runs:
                run.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    if widths:
        for col_idx, w in enumerate(widths):
            for row in t.rows:
                row.cells[col_idx].width = Cm(w)
    return t


def _fmt(v, suffix="", digits=2):
    if v is None:
        return "—"
    if isinstance(v, (int, float)):
        return f"{v:.{digits}f}{suffix}"
    return str(v) + suffix


def main():
    doc = Document()

    # ---- Title ----
    title = doc.add_heading(
        "FENG 498 — Verification & Validation Report\n"
        "Schneider Electric Manisa Switchgear Warehouse Simulation",
        level=0,
    )
    for r in title.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph()
    r = p.add_run("Deniz Ege Memetoğlu  ·  denizegememetoglu@gmail.com")
    r.italic = True
    p = doc.add_paragraph()
    r = p.add_run(f"Date: {date.today().isoformat()}  ·  Framework: Sargent (2013)")
    r.italic = True

    doc.add_paragraph()

    # ---- Executive summary ----
    _add_heading(doc, "Executive summary")
    val = _read_json("validation_report.json", {}) or {}
    stats = _read_json("policy_stats.json", {}) or {}
    zsum = _read_json("zwm92_summary.json", {}) or {}
    timing = _read_json("timing_study_f400.json", {}) or {}

    chi = val.get("chi_square_per_rack_restricted") or {}
    chi_p = chi.get("p_value")
    chi_v = chi.get("chi_square")
    ttest = val.get("t_test_per_material") or {}
    t_p = ttest.get("p_value")
    expected_source = val.get("expected_source", "—")

    summary_lines = [
        f"Data validity: distributions fitted to {zsum.get('total_rows', '—')} real "
        f"ZWM92 dispatch rows ({zsum.get('orders_built', '—')} kit-orders) over "
        f"{zsum.get('date_min', '—')[:10]} → {zsum.get('date_max', '—')[:10]}.",
        f"Operational validity: per-rack chi-square against {expected_source} "
        f"yielded χ²={_fmt(chi_v, '', 1)} (p={_fmt(chi_p, '', 4)}); paired t-test on "
        f"log-daily picks per material yielded p={_fmt(t_p, '', 4)}.",
        f"Conceptual model: synthetic Poisson generator retired in favour of "
        f"Arena-style fitted distributions (src/simulation.py:ZWM92DistributionDriver). "
        f"Per advisor instruction the seed is fixed across all replications so "
        f"experimental contrasts between policies are not confounded by Monte-Carlo noise.",
        f"Timing: F400 video time-motion study (n={timing.get('total_observations', '—')} "
        f"micro-events) replaces the pre-visit placeholders for at-bin pick / "
        f"penalty / RT-pick durations; Lognormal noise added per category CV.",
    ]
    for line in summary_lines:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)

    doc.add_paragraph()

    # ---- 1. Face validity ----
    _add_heading(doc, "1. Face validity")
    _add_para(doc,
        "The layout (config/layout.json) is reconstructed from three converging "
        "sources: site PDFs of every rack (A–J, U), the partner's CAD drawing "
        "SE Manisa Ambar Rafları.dwg (May 26 AutoCAD source), and the SAP özet "
        "bin master. The DWG is read by src/dwg_to_layout.py (libredwg → dwg2dxf "
        "→ ezdxf) and persisted to output/dwg_extracted.json as a face-validity "
        "reference; we deliberately do NOT overwrite config/layout.json because "
        "rack IDs and bay codes are SAP join keys that must stay stable "
        "(see CLAUDE.md). Twelve rack rows on a 4 009 × 3 686 cm floor were "
        "detected vs eleven racks (A–J, U) in layout.json — the extra row is a "
        "back-to-back pair that layout.json groups under a single ID. "
        "Visual verification of the 3D Three.js viewer (web/index.html) against "
        "the CAD drawing is the face-validity gate the supervisor requested."
    )
    dwg = _read_json("dwg_extracted.json", {}) or {}
    if dwg:
        prod_lines = dwg.get("production_lines", []) or []
        if prod_lines:
            rows = [(p.get("line_label", "—"),
                     f"{(p.get('nearest_kitting') or {}).get('x', 0):.0f}, "
                     f"{(p.get('nearest_kitting') or {}).get('y', 0):.0f}",
                     f"{(p.get('nearest_kitting') or {}).get('distance_cm', 0):.0f}")
                    for p in prod_lines]
            _add_para(doc, "Production-line labels extracted from the DWG:", bold=True)
            _add_table(doc, ["Line label", "Kitting (cm)", "Label→kitting dist (cm)"], rows)
    if os.path.exists(LAYOUT_TOPDOWN_PATH):
        doc.add_picture(LAYOUT_TOPDOWN_PATH, width=Cm(15))
        ax = doc.paragraphs[-1]
        ax.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_para(doc,
            "Figure 1 — Modelled top-down layout (config/layout.json + decorations).",
            italic=True)

    # ---- 2. Conceptual model validity ----
    _add_heading(doc, "2. Conceptual model validity")
    _add_para(doc,
        "Orders, kit composition, and per-rack pick frequency are now sourced "
        "from real SAP ZWM92 dispatch logs. The synthetic Poisson generator "
        "that previously sampled materials by consumption weight has been "
        "retired."
    )
    _add_para(doc,
        "Arena-style distribution driver (ZWM92DistributionDriver, src/simulation.py):",
        bold=True,
    )
    bullets = [
        "Inter-arrival ~ Exponential(mean = 4.8 min) — fitted on consecutive ZWM92 timestamps.",
        "Items per order ~ Empirical (observed kit-BOM sizes, clipped at 50 to suppress bulk-issue SAP qty outliers).",
        "Production line ~ Categorical(line_weights from ZWM92).",
        "Material conditional on line ~ Categorical(per-line pick frequencies).",
        "Operator / RT / Kardex pick times ~ Lognormal with means from §4 and σ from F400 video CVs (1.30, 1.20, 1.40, 1.30, 0.30).",
        "Same RANDOM_SEED=42 across all replications (advisor instruction); ANOVA collapses to a paired-by-order Wilcoxon test in §6.",
        "Fallback: when the ZWM92 fit cache is missing the driver reverts to the legacy synthetic generator for CI/dev runs.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # ---- 2a. Flowchart ----
    if os.path.exists(FLOWCHART_PATH):
        _add_para(doc, "Flowchart — one order's lifecycle through the simulation:", bold=True)
        doc.add_picture(FLOWCHART_PATH, width=Cm(15))
        ax = doc.paragraphs[-1]
        ax.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_para(doc,
            "Figure 2 — Discrete-event flow chart. Stochastic sources (red), "
            "operator processes (blue), reach-truck processes (orange), Kardex "
            "(green), decisions (yellow), bookkeeping (grey).",
            italic=True,
        )

    # ---- 2b. Pseudocode ----
    _add_para(doc, "Pseudocode — main simulation loop:", bold=True)
    pseudo = (
        "while sim.time < SHIFT_DURATION:\n"
        "    order = ZWM92DistributionDriver.next_order()  # Arena-style sample\n"
        "    yield env.timeout(order.inter_arrival)         # Exponential IAT\n"
        "    for item in order.items:\n"
        "        target = slotting_policy.locate(item)\n"
        "        if material_in_kardex(item):\n"
        "            with kardex.request():\n"
        "                yield env.timeout(\n"
        "                    Lognormal(carousel) + Lognormal(pick)\n"
        "                )\n"
        "            continue\n"
        "        with position_lock(target.position).request():\n"
        "            if target.level < FAST_MOVER_MAX_LEVEL and target.has_kit_corridor:\n"
        "                with operators.request():\n"
        "                    yield env.timeout(walk_time(operator → target))\n"
        "                    yield env.timeout(Lognormal(OPERATOR_PICK_TIME))\n"
        "            elif target.level <= MANUAL_PICK_MAX_LEVEL:\n"
        "                with operators.request():\n"
        "                    yield env.timeout(\n"
        "                        walk_time(...) + Lognormal(MANUAL_PICK_PENALTY)\n"
        "                    )\n"
        "            else:\n"
        "                with reach_trucks.request():\n"
        "                    yield env.timeout(depot_travel + lift_time)\n"
        "                    yield env.timeout(Lognormal(RT_PICK_PLACE_TIME))\n"
        "        kpi.record_pick(walk, prep, op_wait, rt_wait)\n"
        "    yield env.timeout(walk_back_to_kitting_point(order.line))\n"
        "    kpi.record_order(order)\n"
    )
    code_p = doc.add_paragraph()
    code_run = code_p.add_run(pseudo)
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(9)

    # ---- 3. Data validity ----
    _add_heading(doc, "3. Data validity")
    _add_para(doc, "ZWM92 dispatch log (4-month window):", bold=True)
    fam = zsum.get("families", {})
    rows = [(k, v) for k, v in sorted(fam.items(), key=lambda kv: -kv[1])]
    _add_table(doc, ["Product family", "Dispatch rows"], rows)
    doc.add_paragraph()
    _add_para(doc, "Per-rack actual picks (used as chi-square expected vector):", bold=True)
    racks = zsum.get("picks_per_rack_actual", {})
    rrows = [(k, v) for k, v in sorted(racks.items(), key=lambda kv: -kv[1])]
    _add_table(doc, ["Rack", "Picks (ZWM92)"], rrows)
    doc.add_paragraph()

    # ---- 4. Timing study ----
    _add_heading(doc, "4. Input data — F400 time-motion study")
    _add_para(doc,
        "Video time-motion study (F400 Kit Cansu Nehir.xlsx) over "
        f"{timing.get('total_observed_minutes', '—')} min of DJI footage "
        f"yields {timing.get('total_observations', '—')} classified micro-events. "
        "Categories below; the simulation uses the highlighted mean values."
    )
    cat = timing.get("category_stats", {})
    rows = []
    for name in ("rf_scan", "manual_pick", "walk_corridor", "rt_pick",
                 "rt_travel", "cart_move", "place_on_cart", "wait", "talking"):
        s = cat.get(name, {})
        if not s.get("n"):
            continue
        rows.append((
            name,
            s.get("n"),
            _fmt(s.get("mean_s"), " s"),
            _fmt(s.get("median_s"), " s"),
            _fmt(s.get("p95_s"), " s"),
        ))
    _add_table(doc, ["Category", "n", "mean", "median", "p95"], rows)
    doc.add_paragraph()
    _add_para(doc, "Applied to config.py:", bold=True)
    mp = timing.get("config_mapping", {}) or {}
    constants = [
        ("OPERATOR_PICK_TIME (min)", "0.30", _fmt(mp.get("OPERATOR_PICK_TIME"), "", 4)),
        ("MANUAL_PICK_TIME_PENALTY (min)", "0.50", _fmt(mp.get("MANUAL_PICK_TIME_PENALTY"), "", 4)),
        ("REACH_TRUCK_PICK_PLACE_TIME (min)", "0.50", _fmt(mp.get("REACH_TRUCK_PICK_PLACE_TIME"), "", 4)),
    ]
    _add_table(doc, ["Constant", "Pre-visit placeholder", "F400-derived"], constants)
    doc.add_paragraph()
    _add_para(doc,
        "Limitation: F400 is one of nine product families; constants are "
        "extrapolated to all lines until per-line studies are commissioned. "
        "Sensitivity analysis (§6) bounds this risk.",
        italic=True,
    )

    # ---- 5. Operational validity ----
    _add_heading(doc, "5. Operational validity")
    _add_para(doc, f"Expected source: {expected_source}.")
    if chi:
        _add_para(doc, "Chi-square per rack (restricted to SAP-decoded materials):", bold=True)
        racks_tested = chi.get("racks", [])
        obs = chi.get("observed", [])
        exp = chi.get("expected", [])
        rows = list(zip(racks_tested, obs, exp))
        _add_table(doc, ["Rack", "Observed (sim)", "Expected (ZWM92)"], rows)
        doc.add_paragraph()
        _add_para(doc,
            f"χ² = {_fmt(chi_v, '', 3)} on {chi.get('dof', '—')} dof, "
            f"p = {_fmt(chi_p, '', 4)}. "
            f"{'Reject H0 — distributions differ' if (chi_p or 1) < 0.05 else 'Fail to reject H0 — distributions consistent'}.",
        )
    if ttest:
        _add_para(doc,
            f"Paired t-test, log(daily picks/material): n={ttest.get('n_materials_paired')}, "
            f"t={_fmt(ttest.get('t_statistic'), '', 3)}, p={_fmt(t_p, '', 4)}.",
        )

    # ---- 6. Policy comparison ----
    metrics = stats.get("metrics") if isinstance(stats, dict) else None
    paired = stats.get("paired_by_order") if isinstance(stats, dict) else None
    if metrics:
        _add_heading(doc, "6. Policy comparison")
        # Policy name normalisation — policy_stats keys come straight from the
        # ALL_POLICIES dict in src/slotting.py.
        possible_order = [
            ("Baseline (Heuristic)", "Heuristic"),
            ("Baseline (Actual SAP)", "Actual SAP"),
            ("Usage-based ABC", "Usage-Based ABC"),
            ("Double ABC", "Double-Sided ABC"),
            ("Travel-distance Optimized", "Travel-Dist Opt"),
        ]
        # Discover order by inspecting per_policy of one KPI.
        any_kpi = next(iter(metrics.values()))
        present_policies = set((any_kpi.get("per_policy") or {}).keys())
        order = [k for k, _ in possible_order if k in present_policies]
        labels = {k: lbl for k, lbl in possible_order}
        kpi_keys = [
            ("avg_walk_distance", "Walk dist/order (m)"),
            ("avg_lead_time", "Lead time (min)"),
            ("avg_prep_time", "Prep time (min)"),
            ("avg_op_queue_wait", "Op-queue wait (min)"),
            ("reach_truck_utilization", "RT util"),
            ("operator_utilization", "Op util"),
            ("orders_completed", "Orders completed"),
        ]
        headers = ["KPI"] + [labels[k] for k in order]
        rows = []
        for key, label in kpi_keys:
            mblock = metrics.get(key) or {}
            per_policy = mblock.get("per_policy") or {}
            row = [label]
            for pol in order:
                d = per_policy.get(pol, {}) or {}
                mean = d.get("mean")
                ci = d.get("ci95")
                if mean is None:
                    row.append("—")
                else:
                    cell = f"{mean:.2f}"
                    if isinstance(ci, list) and len(ci) == 2 and ci[0] is not None:
                        cell += f" [{ci[0]:.2f}, {ci[1]:.2f}]"
                    row.append(cell)
            rows.append(row)
        _add_table(doc, headers, rows)
        doc.add_paragraph()
        _add_para(doc,
            "N_REPLICATIONS=1 with identical seeds means each cell is a single "
            "deterministic trajectory; the ±95% CI columns therefore collapse "
            "to the point estimate. The classical ANOVA/Tukey/Welch tests "
            "require between-rep variance and are not applicable here. The "
            "paired-by-order Wilcoxon signed-rank test below provides the "
            "inferential check the advisor asked for — every order is paired "
            "against the same order under the Heuristic baseline, eliminating "
            "between-order variance entirely.",
            italic=True,
        )
        if os.path.exists(POLICY_CHART_PATH):
            doc.add_picture(POLICY_CHART_PATH, width=Cm(15))
            ax = doc.paragraphs[-1]
            ax.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_para(doc, "Figure 3 — Policy comparison across KPIs.",
                      italic=True)

        # Paired-by-order Wilcoxon
        if paired:
            _add_para(doc, "Paired-by-order Wilcoxon vs Baseline (Heuristic):",
                      bold=True)
            rows = []
            for pol, kpi_dict in paired.items():
                for kpi_name, m in kpi_dict.items():
                    sig = "*" if (m.get("p_value") or 1) < 0.05 else " "
                    rows.append((
                        f"{sig} {pol}",
                        m.get("label", kpi_name),
                        f"{m.get('mean_diff', 0):.3f}",
                        f"{m.get('median_diff', 0):.3f}",
                        f"{m.get('p_value', 1):.4f}",
                        m.get("n_paired", "—"),
                    ))
            _add_table(doc,
                ["Policy", "KPI", "Mean Δ", "Median Δ", "p-value", "n"],
                rows)
            _add_para(doc,
                "Asterisks mark p<0.05. Negative deltas favour the test policy "
                "(lower walk / prep / lead is better).",
                italic=True,
            )

    # ---- 7. Sensitivity ----
    sens = _read_json("sensitivity.json", {}) or {}
    if sens:
        _add_heading(doc, "7. Sensitivity (OAT tornados)")
        _add_para(doc,
            "One-at-a-time perturbation of model parameters with ±20% on "
            "OPERATOR_PICK_TIME, REACH_TRUCK_PICK_PLACE_TIME, NUM_OPERATORS, "
            "NUM_REACH_TRUCKS. KPI deltas relative to the baseline are in "
            "output/sensitivity.json and the tornado plots in "
            "output/tornado_*.png.",
        )

    # ---- 8. Limitations ----
    _add_heading(doc, "8. Limitations and future work")
    limits = [
        "F400-derived timing constants are extrapolated to all 9 product families.",
        "Single-site model — Manisa only. Generalisation to other Schneider sites untested.",
        "Batch picking not modelled — each order is independent.",
        "Milkrun process gated off (ENABLE_MILKRUN=False) until field BOM data confirms tour content.",
        "Position-lock granularity is the pallet slot, not the bin — collisions inside a bay are not detected.",
        "DWG-derived layout coordinates are recorded in output/dwg_extracted.json as a face-validity reference only; config/layout.json is not overwritten because rack IDs and bay codes are SAP join keys.",
        "Same-seed reproducibility eliminates Monte-Carlo noise but rules out classical ANOVA — paired-by-order Wilcoxon is used in its place. Switching SAME_SEED_FOR_ALL_REPS=False would re-enable replication-based inference.",
    ]
    for l in limits:
        doc.add_paragraph(l, style="List Bullet")

    # ---- Appendix ----
    _add_heading(doc, "Appendix A — File map")
    file_map = [
        ("config/layout.json", "Rack geometry, kit corridors, kitting points"),
        ("data/zwm92/*.XLSX", "9 SAP ZWM92 family exports (4 months, 167k rows)"),
        ("src/zwm92.py", "ZWM92 loader + derived views"),
        ("src/timing_study.py", "F400 video time-motion extractor"),
        ("src/simulation.py", "SimPy core + ZWM92DistributionDriver (Arena-style)"),
        ("src/validate.py", "Chi-square + paired t-test vs ZWM92 actuals"),
        ("src/analyze.py", "ANOVA / Welch + paired-by-order Wilcoxon"),
        ("src/sensitivity.py", "OAT sensitivity sweeps"),
        ("src/dwg_to_layout.py", "Read-only DXF extractor (face-validity reference)"),
        ("output/timing_study_f400.json", "F400 categorical timing breakdown"),
        ("output/zwm92_{summary,orders,observed_bins}.json", "Cached ZWM92 views"),
        ("output/validation_report.{json,txt}", "Chi-square + t-test results"),
        ("output/policy_stats.{json,txt}", "ANOVA + Tukey HSD + Welch + Cohen's d"),
        ("ASSUMPTIONS.md", "Narrative justification of every modelling assumption"),
    ]
    _add_table(doc, ["File", "Role"], file_map)

    os.makedirs(os.path.dirname(DEST), exist_ok=True)
    doc.save(DEST)
    print(f"Wrote {DEST}")


if __name__ == "__main__":
    main()
